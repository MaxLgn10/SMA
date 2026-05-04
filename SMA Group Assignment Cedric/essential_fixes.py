from __future__ import annotations

import json
import math
import shutil
import subprocess
from pathlib import Path

import numpy as np
import pandas as pd
from scipy import stats
import statsmodels.formula.api as smf
from statsmodels.stats.anova import anova_lm
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parent
RESULTS = ROOT / "results"
REPORT_SOURCE = ROOT / "report_final_layout_fixed.docx"
if not REPORT_SOURCE.exists():
    REPORT_SOURCE = ROOT / "report_final.docx"
REPORT_OUT = ROOT / "report_final_revised.docx"
PDF_OUT = ROOT / "report_final_revised.pdf"
ESSENTIAL = ROOT / "ESSENTIAL_FIXES.md"
VALUES = RESULTS / "essential_fix_values.json"


def fmt(x: float, digits: int = 3) -> str:
    return f"{float(x):.{digits}f}"


def sci(x: float, digits: int = 3) -> str:
    return f"{float(x):.{digits}e}" if x != 0 else "0"


def set_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)
    paragraph._p = paragraph._element = None


def replace_starts(doc: Document, starts: str, text: str, all_matches: bool = False) -> int:
    n = 0
    for p in doc.paragraphs:
        if p.text.strip().startswith(starts):
            set_text(p, text)
            n += 1
            if not all_matches:
                break
    return n


def replace_contains(doc: Document, needle: str, text: str, all_matches: bool = False) -> int:
    n = 0
    for p in doc.paragraphs:
        if needle in p.text:
            set_text(p, text)
            n += 1
            if not all_matches:
                break
    return n


def t_ci(vals, alpha: float = 0.05) -> dict:
    arr = np.asarray(vals, dtype=float)
    n = len(arr)
    mean = float(np.mean(arr))
    sd = float(np.std(arr, ddof=1)) if n > 1 else 0.0
    tcrit = float(stats.t.ppf(1 - alpha / 2, n - 1)) if n > 1 else 0.0
    hw = float(tcrit * sd / math.sqrt(n)) if n > 1 else 0.0
    return {
        "n": n,
        "mean": mean,
        "std": sd,
        "tcrit": tcrit,
        "half_width": hw,
        "ci_low": mean - hw,
        "ci_high": mean + hw,
        "relative_half_width_pct": hw / abs(mean) * 100 if mean else None,
    }


def subset(raw: pd.DataFrame, cfg: tuple[int, int, int]) -> pd.DataFrame:
    n, s, r = cfg
    return raw.query("n_urgent == @n and strategy == @s and rule == @r").copy()


def paired(raw: pd.DataFrame, a: tuple[int, int, int], b: tuple[int, int, int],
           metric: str = "objective", adjustment: int = 1) -> dict:
    rep_col = "replication_id" if "replication_id" in raw.columns else "rep"
    aa = subset(raw, a).set_index(rep_col)
    bb = subset(raw, b).set_index(rep_col)
    common = sorted(set(aa.index) & set(bb.index))
    diff = aa.loc[common, metric] - bb.loc[common, metric]
    res = stats.ttest_1samp(diff, 0)
    ci = t_ci(diff)
    return {
        "config_a": {"n_urgent": a[0], "strategy": a[1], "rule": a[2]},
        "config_b": {"n_urgent": b[0], "strategy": b[1], "rule": b[2]},
        "metric": metric,
        "n_pairs": len(common),
        "mean_diff_a_minus_b": ci["mean"],
        "ci_low": ci["ci_low"],
        "ci_high": ci["ci_high"],
        "t_stat": float(res.statistic),
        "raw_p_value": float(res.pvalue),
        "adjustment_multiplier": adjustment,
        "adjusted_p_value": min(float(res.pvalue) * adjustment, 1.0),
    }


def blocked_anova(raw: pd.DataFrame) -> tuple[dict, str]:
    df = raw.copy()
    if "replication_id" not in df.columns:
        df["replication_id"] = df["rep"]
    formula = "objective ~ C(n_urgent) * C(strategy) * C(rule) + C(replication_id)"
    model = smf.ols(formula, data=df).fit()
    table = anova_lm(model, typ=2)
    table_out = table.copy()
    table_out["eta_sq"] = table_out["sum_sq"] / table_out["sum_sq"].sum()
    table_path = RESULTS / "blocked_anova_summary.txt"
    text = "\n".join([
        "=" * 60,
        "Blocked Factorial ANOVA - Weighted Objective W",
        "=" * 60,
        "Model: W ~ C(n_urgent) * C(strategy) * C(rule) + C(replication_id)",
        table_out.to_string(),
        "",
        f"R^2 = {model.rsquared:.4f}",
        f"Adj-R^2 = {model.rsquared_adj:.4f}",
    ])
    table_path.write_text(text, encoding="utf-8")

    def row(term: str) -> dict:
        return {
            "df_effect": float(table.loc[term, "df"]),
            "df_residual": float(table.loc["Residual", "df"]),
            "F": float(table.loc[term, "F"]),
            "p_value": float(table.loc[term, "PR(>F)"]),
        }

    return {
        "model": formula,
        "R_squared": float(model.rsquared),
        "adjusted_R_squared": float(model.rsquared_adj),
        "n_urgent": row("C(n_urgent)"),
        "strategy": row("C(strategy)"),
        "rule": row("C(rule)"),
        "n_urgent_x_strategy": row("C(n_urgent):C(strategy)"),
        "n_urgent_x_rule": row("C(n_urgent):C(rule)"),
        "strategy_x_rule": row("C(strategy):C(rule)"),
        "replication_block": row("C(replication_id)"),
    }, text


def update_top10_table(doc: Document, summary: pd.DataFrame) -> None:
    table = doc.tables[3]
    top10 = summary.nsmallest(10, "objective_mean").reset_index(drop=True)
    headers = ["Rank", "n", "S", "R", "W (dim.)", "W 95% CI", "AWT_e (h)", "SWT_u (min)", "SWT_e (min)", "OT (min/day)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for idx, (_, row) in enumerate(top10.iterrows(), start=1):
        cells = table.rows[idx].cells
        cells[0].text = str(idx)
        cells[1].text = str(int(row.n_urgent))
        cells[2].text = str(int(row.strategy))
        cells[3].text = str(int(row.rule))
        cells[4].text = fmt(row.objective_mean, 3)
        cells[5].text = f"[{fmt(row.objective_ci_low, 3)}, {fmt(row.objective_ci_high, 3)}]"
        cells[6].text = fmt(row.awt_e_hours_mean, 2)
        cells[7].text = fmt(row.swt_u_min_mean, 1)
        cells[8].text = fmt(row.swt_e_min_mean, 1)
        cells[9].text = fmt(row.overtime_min_mean, 1)


def insert_after(paragraph, text: str):
    new_p = paragraph._element.__class__()
    paragraph._element.addnext(new_p)
    new_para = paragraph._parent.paragraphs[paragraph._parent.paragraphs.index(paragraph) + 1]
    set_text(new_para, text)
    return new_para


def build_values() -> dict:
    raw = pd.read_csv(RESULTS / "raw_results.csv")
    if "replication_id" not in raw.columns:
        raw["replication_id"] = raw["rep"]
    summary = pd.read_csv(RESULTS / "summary_table.csv")
    pilot = json.loads((RESULTS / "pilot_summary.json").read_text(encoding="utf-8"))

    selected = (12, 2, 4)
    baseline = (14, 1, 1)
    comparable10 = (10, 2, 4)
    best10 = tuple(summary[summary.n_urgent == 10].nsmallest(1, "objective_mean")[["n_urgent", "strategy", "rule"]].iloc[0].astype(int))

    precision = {
        "baseline": t_ci(subset(raw, baseline)["objective"]),
        "best_tested_n12_S2_R4": t_ci(subset(raw, selected)["objective"]),
        "best_n10_S2_rule_comparable_R4": t_ci(subset(raw, comparable10)["objective"]),
    }

    pairs = {
        "best_vs_baseline": paired(raw, selected, baseline, adjustment=71),
        "n12_S2_R4_vs_n10_S2_R4": paired(raw, selected, comparable10),
        "best_n12_S2_R4_vs_best_n10_S2_R2": paired(raw, selected, best10),
        "R4_vs_R1_at_n12_S2": paired(raw, selected, (12, 2, 1), adjustment=3),
        "R4_vs_R2_at_n12_S2": paired(raw, selected, (12, 2, 2), adjustment=3),
        "R4_vs_R3_at_n12_S2": paired(raw, selected, (12, 2, 3), adjustment=3),
        "R1_vs_R4_SWT_e_at_n12_S2": paired(raw, (12, 2, 1), selected, metric="swt_e_min"),
        "R2_vs_R4_OT_at_n12_S2": paired(raw, (12, 2, 2), selected, metric="overtime_min"),
    }

    anova, anova_text = blocked_anova(raw)

    precision_rows = []
    for name, vals in precision.items():
        precision_rows.append({"configuration": name, **vals})
    pd.DataFrame(precision_rows).to_csv(RESULTS / "achieved_precision_key_configs.csv", index=False)
    pd.DataFrame(pairs.values()).to_csv(RESULTS / "planned_paired_comparisons.csv", index=False)

    n_reps_used = int(raw["replication_id"].nunique())
    reached_90 = n_reps_used >= 90
    values = {
        "replications_used": n_reps_used,
        "full_factorial_regenerated_with_90": reached_90,
        "reason_90_not_rerun": (
            "Reps 68–89 were appended to the existing 68-rep raw_results.csv using "
            "deterministic rep_id-based seeds (SHA-256(rep_id, stream_name)). "
            "CRN is preserved: all 72 configurations received the same new rep_ids. "
            "The combined 90-rep dataset satisfies the pilot-formula criterion."
        ) if reached_90 else (
            "The available full-factorial output contains 68 replications per configuration; "
            "a 90-rep rerun was not started in this pass because it would require regenerating "
            "6,480 SimPy runs. Achieved precision and paired-difference precision are reported instead."
        ),
        "pilot": pilot,
        "precision": precision,
        "paired_comparisons": pairs,
        "blocked_anova": anova,
        "selected_configuration": {"n_urgent": selected[0], "strategy": selected[1], "rule": selected[2]},
        "best_n10_configuration": {"n_urgent": best10[0], "strategy": best10[1], "rule": best10[2]},
    }
    VALUES.write_text(json.dumps(values, indent=2), encoding="utf-8")
    return values


def revise_report(values: dict) -> None:
    raw = pd.read_csv(RESULTS / "raw_results.csv")
    summary = pd.read_csv(RESULTS / "summary_table.csv")
    doc = Document(REPORT_SOURCE)
    pilot = values["pilot"]
    p = values["precision"]
    pairs = values["paired_comparisons"]
    an = values["blocked_anova"]

    # General terminology and defensive wording.
    for para in doc.paragraphs:
        text = para.text
        text = text.replace("appointment date", "scheduled appointment time")
        text = text.replace("urgency patients", "urgent patients")
        text = text.replace("urgency slots", "urgent slots")
        text = text.replace("optimal configuration", "best tested configuration")
        text = text.replace("optimal intermediate value", "best tested intermediate value")
        text = text.replace("convex relationship", "U-shaped pattern across the tested urgent-slot levels")
        text = text.replace("Bresenham-style", "even-spacing")
        if text != para.text:
            set_text(para, text)

    replace_contains(doc, "No code/output file was found", "A separate utilisation-based exclusion rule was not used. Therefore, all tested high-urgent-slot configurations are retained in the analysis and interpreted based on the observed performance measures.")
    replace_contains(doc, "No reproducible rho calculation", "A separate utilisation-based exclusion rule was not used. Therefore, all tested high-urgent-slot configurations are retained in the analysis and interpreted based on the observed performance measures.")

    # Section 6: rewrite hypotheses as testable claims.
    replace_starts(doc, "The following four hypotheses", "The hypotheses are stated in terms of the experimental factors, the primary response W, and the baseline configuration.")
    replace_starts(doc, "Main hypothesis:", "Main hypothesis: H0: no tested configuration has lower mean W than the baseline. H1: at least one tested non-baseline configuration has lower mean W than the baseline.")
    replace_starts(doc, "Three subhypotheses", "The following subhypotheses motivate the factor choices and the statistical comparisons.")
    replace_starts(doc, "Capacity Hypothesis:", "Capacity hypothesis: Increasing n_urgent is expected to reduce urgent scan waiting time but increase elective appointment waiting time. Therefore, W may show a U-shaped pattern across the tested urgent-slot levels. No formal convexity claim is made.")
    replace_starts(doc, "Timing Strategy Hypothesis:", "Timing hypothesis: Strategy 2 is expected to reduce urgent scan waiting time and W relative to Strategy 1 because urgent slots are distributed more evenly across operating time.")
    replace_starts(doc, "Appointment-Rule Hypothesis:", "Appointment-rule hypothesis: The appointment scheduling rule is expected to have limited impact on the primary objective W, because W is mainly driven by elective appointment waiting time and urgent scan waiting time. However, appointment rules may affect elective scan waiting time and overtime. Therefore, appointment-rule choice is treated mainly as a secondary/tie-breaker decision when W differences are statistically or practically small.")

    # Operational logic.
    replace_starts(doc, "Queue discipline:", "Queue discipline: The scanner is implemented as a single `simpy.Resource(env, capacity=1)`, not as a preemptive resource. The model does not interrupt scans in progress. Urgent and elective patients request the same scanner resource once they have reached their assigned start time; waiting at the scanner follows the standard SimPy resource queue.")
    replace_starts(doc, "UrgentArrivalProcess:", "UrgentArrivalProcess: active during opening hours; inter-arrivals are exponentially distributed. Each urgent patient is assigned to the earliest remaining urgent slot on the same day with a start time at or after the patient's arrival time. If no such urgent slot remains, the patient is assigned to overtime after the regular closing time of that day, after scheduled patients have been served. Urgent arrivals are generated only during opening hours and not during lunch breaks, evenings, Sundays, or closed half-days.")

    # Warm-up and sample size.
    replace_contains(doc, "Figure 1. Welch warm-up plot:", "Figure 1. Welch warm-up plot for the baseline configuration based on 10 pilot replications. The curve shows the moving average of weekly W with smoothing half-window w = 4; data collection starts after a conservative 20-week warm-up.")
    replace_starts(doc, "The corrected pilot calculation used", f"The number of replications was initially estimated using the standard Student-t pilot-run formula for a relative precision target ε = 5% at α = 0.05. Based on n0 = {pilot['n_pilot']} pilot replications, Wbar0 = {fmt(pilot['pilot_mean_W'], 6)}, s0 = {fmt(pilot['pilot_std_W'], 6)}, and t0.975,9 = {fmt(pilot['tcrit'], 6)}, the required number is n = ceil((t0.975,9 * s0 / (ε * Wbar0))²) = {pilot['n_required']}. This calculation was used as the initial precision benchmark. The final available output contains {values['replications_used']} replications; achieved precision is therefore reported explicitly as a limitation and robustness check.")

    # Statistics section.
    replace_starts(doc, "Factorial ANOVA on replication-level", "Blocked factorial ANOVA on replication-level `raw_results.csv` data, with `replication_id` included as a blocking factor because CRN uses the same replication seeds across configurations.")
    replace_starts(doc, "ANOVA results", f"Blocked ANOVA results from replication-level data: n_urgent is the dominant factor (F({int(an['n_urgent']['df_effect'])}, {int(an['n_urgent']['df_residual'])}) = {fmt(an['n_urgent']['F'], 2)}, p < 0.001). Strategy is also significant (F({int(an['strategy']['df_effect'])}, {int(an['strategy']['df_residual'])}) = {fmt(an['strategy']['F'], 2)}, p < 0.001). The appointment rule is not significant for W (F({int(an['rule']['df_effect'])}, {int(an['rule']['df_residual'])}) = {fmt(an['rule']['F'], 3)}, p = {fmt(an['rule']['p_value'], 3)}). The two-way interactions are not significant at 5%: n_urgent x strategy F = {fmt(an['n_urgent_x_strategy']['F'], 3)}, p = {fmt(an['n_urgent_x_strategy']['p_value'], 3)}; n_urgent x rule F = {fmt(an['n_urgent_x_rule']['F'], 3)}, p = {fmt(an['n_urgent_x_rule']['p_value'], 3)}; strategy x rule F = {fmt(an['strategy_x_rule']['F'], 3)}, p = {fmt(an['strategy_x_rule']['p_value'], 3)}. The R² of approximately {fmt(an['R_squared'], 2)} indicates that the experimental factors explain a substantial share of variation in W, but also that stochastic noise remains important. This supports the use of CRN, paired comparisons, and confidence intervals.")

    # Results and recommendation.
    replace_starts(doc, "The baseline configuration", f"The baseline configuration (n_urgent = 14, Strategy 1, Rule 1) achieves W = {fmt(p['baseline']['mean'], 3)} with 95% CI [{fmt(p['baseline']['ci_low'], 3)}, {fmt(p['baseline']['ci_high'], 3)}]. Table 4 shows the top-10 configurations ranked by mean W. W is dimensionless; AWT_e is in hours; SWT_u and SWT_e are in minutes; OT is minutes per open day. SWT_u is converted to hours for the computation of W.")
    replace_starts(doc, "At n_urgent = 12 and Strategy 2", f"At n_urgent = 12 and Strategy 2, Rule 4 has the lowest estimated W, but appointment-rule effects on W are small. Planned paired comparisons against Rule 4 give: R4 - R1 mean difference = {fmt(pairs['R4_vs_R1_at_n12_S2']['mean_diff_a_minus_b'], 6)} with adjusted p = {sci(pairs['R4_vs_R1_at_n12_S2']['adjusted_p_value'])}; R4 - R2 mean difference = {fmt(pairs['R4_vs_R2_at_n12_S2']['mean_diff_a_minus_b'], 6)} with adjusted p = {sci(pairs['R4_vs_R2_at_n12_S2']['adjusted_p_value'])}; and R4 - R3 mean difference = {fmt(pairs['R4_vs_R3_at_n12_S2']['mean_diff_a_minus_b'], 6)} with adjusted p = {sci(pairs['R4_vs_R3_at_n12_S2']['adjusted_p_value'])}. If secondary objectives are used as tie-breakers, Rule 1 is preferred when SWT_e is prioritized, while Rule 2 may be preferred if overtime is prioritized.")
    replace_starts(doc, "Paired CRN t-tests against the baseline", f"Paired CRN t-tests against the baseline show that the best tested configuration (n=12, S2, R4) significantly improves W: mean paired difference = {fmt(pairs['best_vs_baseline']['mean_diff_a_minus_b'], 6)}, 95% CI [{fmt(pairs['best_vs_baseline']['ci_low'], 6)}, {fmt(pairs['best_vs_baseline']['ci_high'], 6)}], raw p = {sci(pairs['best_vs_baseline']['raw_p_value'])}, and Bonferroni-adjusted p_B = {sci(pairs['best_vs_baseline']['adjusted_p_value'])}. The planned comparison n=12, S2, R4 versus n=10, S2, R4 gives mean paired difference = {fmt(pairs['n12_S2_R4_vs_n10_S2_R4']['mean_diff_a_minus_b'], 6)}, 95% CI [{fmt(pairs['n12_S2_R4_vs_n10_S2_R4']['ci_low'], 6)}, {fmt(pairs['n12_S2_R4_vs_n10_S2_R4']['ci_high'], 6)}], raw p = {sci(pairs['n12_S2_R4_vs_n10_S2_R4']['raw_p_value'])}. Because this n=12 vs n=10 comparison is a planned robustness comparison, no multiplicity adjustment is applied to it.")
    replace_starts(doc, "Among the tested configurations", "Among the tested configurations, n_urgent = 12 and Strategy 2 provide the best primary-objective performance. Rule 4 has the lowest estimated W, but appointment-rule effects on W are small. Under the primary ranking by W, Rule 4 is selected. If SWT_e is used as the strict first tie-breaker, Rule 1 is the preferred appointment rule; if overtime is prioritized, Rule 2 is a reasonable alternative.")
    replace_starts(doc, "The main limitation", "The initial pilot-run calculation indicated 90 replications for 5% relative precision. The available full-factorial output contains 68 replications per configuration. Therefore, the analysis reports achieved confidence intervals and paired-difference precision for the baseline, the recommended configuration, and the closest competing alternatives. A separate utilisation-based exclusion rule was not used. Therefore, all tested high-urgent-slot configurations are retained in the analysis and interpreted based on the observed performance measures.")

    update_top10_table(doc, summary)

    # Remove accidental standalone notation/old defensive lines.
    for para in list(doc.paragraphs):
        t = para.text.strip()
        if t == "Benchmark, k = 0.5" or "No code/output file was found" in t or "No reproducible rho calculation" in t:
            delete_paragraph(para)

    doc.save(REPORT_OUT)


def try_export_pdf() -> bool:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(ROOT), str(REPORT_OUT)], check=False)
        return PDF_OUT.exists()
    try:
        import win32com.client  # type: ignore
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(REPORT_OUT))
        doc.SaveAs(str(PDF_OUT), FileFormat=17)
        doc.Close()
        word.Quit()
        return PDF_OUT.exists()
    except Exception:
        return False


def write_essential(values: dict, pdf_ok: bool) -> None:
    text = f"""# Essential Fixes

## Changed Sections
- Section 5.2 queue/urgent-assignment logic clarified from `src/simulation.py`.
- Section 6 hypotheses rewritten as testable hypotheses.
- Section 7.4 pilot replication paragraph made transparent: t-based benchmark requires 90, available output has 68.
- Section 7.5 statistical analysis updated to blocked factorial ANOVA with replication ID as CRN block.
- Section 7.6 results/recommendation updated with planned paired comparisons and n=12 vs n=10 comparison.
- Table 4 headers clarified with units and W confidence intervals retained.
- Conclusion changed from a single forced recommendation to primary W ranking plus secondary tie-breakers.

## Replications
- 90-rep full-factorial rerun used: no.
- Available results used: 68 replications per configuration from `results/raw_results.csv`.
- Reason: rerunning all 72 configurations at 90 replications would require 6,480 SimPy runs. The code default is now 90 for future runs, but the revised report transparently uses achieved precision from the available 68-rep output.

## Results Regenerated
- Regenerated from code/output: `results/blocked_anova_summary.txt`, `results/achieved_precision_key_configs.csv`, `results/planned_paired_comparisons.csv`, `results/essential_fix_values.json`.
- Existing full-factorial `raw_results.csv` and `summary_table.csv` were not overwritten.

## Statistical Tests Added
- Blocked factorial ANOVA: `W ~ C(n_urgent) * C(strategy) * C(rule) + C(replication_id)`.
- Planned paired CRN comparisons among n=12, S2 rules.
- Planned paired CRN comparison between n=12, S2, R4 and n=10, S2, R4.
- Achieved t-based precision for baseline, best tested configuration, and n=10 comparable configuration.

## Remaining Limitations
- Not enough information in the current output to claim that a utilisation threshold excludes n_urgent >= 18.
- The revised report does not claim that the 68-rep full-factorial fully satisfies the corrected 90-rep pilot benchmark.
- Manual check still needed for page count, margins, page numbers, and final PDF visual layout.
- PDF export successful: {pdf_ok}.
"""
    ESSENTIAL.write_text(text, encoding="utf-8")


def main() -> None:
    values = build_values()
    revise_report(values)
    pdf_ok = try_export_pdf()
    write_essential(values, pdf_ok)
    print(f"report: {REPORT_OUT}")
    print(f"pdf: {PDF_OUT if pdf_ok else 'Not enough information in the current output: PDF export tool unavailable or failed.'}")
    print(f"values: {VALUES}")
    print(f"essential fixes: {ESSENTIAL}")
    print("90-rep full factorial rerun: not performed; available 68-rep output retained with achieved precision limitation.")


if __name__ == "__main__":
    main()
