#!/usr/bin/env python3
"""Generate the SMA Project Assignment report in .docx format matching the previous exercises layout."""
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── colour palette (matched from Report_Final_Version.docx) ─────────────────
BLUE     = RGBColor(0x1E, 0x64, 0xC8)   # heading blue
BLUE_LT  = RGBColor(0x4C, 0x94, 0xD8)   # caption blue
BLACK    = RGBColor(0x00, 0x00, 0x00)

RESULTS = r"C:\Users\Cedri\Documents\2.HIR\Simulation\github\SMA\SMA Group Assignment Cedric\results"
OUT     = r"C:\Users\Cedri\Documents\2.HIR\Simulation\github\SMA\SMA Group Assignment Cedric\report_final_revised_v6.docx"

# ── helpers ──────────────────────────────────────────────────────────────────
def _run(p, text, bold=False, italic=False, underline=False,
         color=BLACK, size_pt=10):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.underline = underline
    r.font.size = Pt(size_pt)
    r.font.color.rgb = color
    return r

def _para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    return p

def title_main(doc, text):
    p = _para(doc, before=0, after=4)
    _run(p, text, bold=True, color=BLUE, size_pt=30)

def title_sub(doc, text):
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)
    _run(p, text, color=BLUE, size_pt=20)

def heading1(doc, text):
    p = _para(doc, before=14, after=4)
    _run(p, text, bold=True, color=BLUE, size_pt=12)

def heading2(doc, text):
    p = _para(doc, before=8, after=3)
    _run(p, text, bold=True, color=BLUE, size_pt=10)

def heading_appendix(doc, text):
    p = _para(doc, before=10, after=4)
    _run(p, text, bold=True, color=BLUE, size_pt=13)

def body(doc, text):
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=4)
    _run(p, text, color=BLACK, size_pt=10)
    return p

def body_para(doc):
    """Return an empty justified paragraph so the caller can add mixed runs."""
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=4)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        _run(p, bold_prefix, bold=True, color=BLACK, size_pt=10)
        _run(p, text, color=BLACK, size_pt=10)
    else:
        _run(p, text, color=BLACK, size_pt=10)

def caption(doc, text):
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=10)
    _run(p, text, italic=True, color=BLUE_LT, size_pt=8)

def fig(doc, fname, w=5.5, cap=None):
    path = os.path.join(RESULTS, fname)
    if os.path.exists(path):
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2)
        p.add_run().add_picture(path, width=Inches(w))
    if cap:
        caption(doc, cap)

def tbl_header_cell(cell, text, size=9):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size)

def tbl_cell(cell, text, size=9):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(size)

def set_margins(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.2)
        s.page_width  = Cm(21.6)
        s.page_height = Cm(27.9)

# ── document ─────────────────────────────────────────────────────────────────
def build():
    doc = Document()
    set_margins(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ══ TITLE PAGE ══════════════════════════════════════════════════════════
    title_main(doc, "SIMULATION MODELLING AND ANALYSIS (F000941)")
    title_sub(doc,  "PROJECT ASSIGNMENT – REORGANISATION OF A RADIOLOGY DEPARTMENT")

    for _ in range(3): doc.add_paragraph()

    # TODO: fill in student number and group number before submission
    for name in ["Cédric Claerbout  –  [student number]  |  Group [group number]"]:
        p = _para(doc, before=0, after=2)
        _run(p, name, color=BLACK, size_pt=10)

    for _ in range(6): doc.add_paragraph()

    p = _para(doc, before=0, after=2)
    _run(p, "Prof. Broos Maenhout", bold=True, color=BLACK, size_pt=10)
    p = _para(doc, before=0, after=2)
    _run(p, "Academic Year 2025-2026", bold=True, color=BLACK, size_pt=10)

    doc.add_page_break()

    # ══ TABLE OF CONTENTS ════════════════════════════════════════════════════
    p = _para(doc, before=0, after=8)
    _run(p, "TABLE OF CONTENTS", bold=True, color=BLUE, size_pt=14)

    toc = [
        ("1.", "Introduction: the objectives of the experiment", False),
        ("2.", "Operational variability of the input data", False),
        ("3.", "Evaluation output measure (responses)", False),
        ("4.", "Design parameters and values under consideration (factors)", False),
        ("5.", "The simulation design", False),
        ("",   "5.1  State descriptor", True),
        ("",   "5.2  Events", True),
        ("",   "5.3  Agenda", True),
        ("",   "5.4  Main routine", True),
        ("",   "5.5  Simulation time", True),
        ("6.", "Hypotheses", False),
        ("7.", "Choice of experimental design", False),
        ("",   "7.1  Full factorial design", True),
        ("",   "7.2  Variance reduction techniques (CRN)", True),
        ("",   "7.3  Warm-up period", True),
        ("",   "7.4  Number of simulation runs", True),
        ("",   "7.5  Statistical analysis", True),
        ("",   "7.6  Results", True),
        ("8.", "Conclusion", False),
        ("",   "References", False),
        ("",   "Appendix", False),
    ]
    for num, title, sub in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        indent = "     " if sub else ""
        _run(p, f"{indent}{num}  {title}".strip(), bold=not sub, color=BLACK, size_pt=10)

    doc.add_page_break()

    # ══ 1. INTRODUCTION ══════════════════════════════════════════════════════
    heading1(doc, "1.  Introduction: the objectives of the experiment")

    body(doc,
        "Hospitals face growing pressure to deliver timely diagnostic services while operating "
        "within constrained equipment capacity. An outpatient radiology department that runs a "
        "single MRI scanner must simultaneously serve two fundamentally different patient "
        "streams: elective patients who book appointments in advance and whose primary concern "
        "is a short waiting time until their appointment date, and urgent patients who arrive "
        "without an appointment and require same-day scanning. The scheduling of outpatient "
        "appointments is a well-studied problem in operations research: Klassen and Rohleder "
        "(1996) demonstrate that the choice of appointment scheduling rule and capacity "
        "allocation significantly affects both patient waiting times and resource utilisation. "
        "Balancing these competing demands requires careful decisions about slot capacity "
        "allocation, the placement of reserved urgent slots within the weekly schedule, and "
        "the rule used to translate slot assignments into appointment times for elective patients.")

    body(doc,
        "The department operates on a cyclic one-week schedule that repeats every 52 weeks. "
        "It is open Monday through Saturday (Sunday closed). Thursday and Saturday are half-days "
        "(morning only, 08:00–12:00). Full days run from 08:00 to 12:00 and 13:00 to 17:00 "
        "with a one-hour lunch break. Each time slot is 15 minutes; the total weekly capacity "
        "is 160 slots. The current configuration allocates 146 elective and 14 urgent slots "
        "per week.")

    body(doc,
        "The central objective is to minimise the weighted sum of patient waiting times:")

    p = body_para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Minimise   W  =  w", color=BLACK, size_pt=10)
    r = p.add_run("e"); r.font.subscript = True; r.font.size = Pt(9)
    _run(p, " × AWT", color=BLACK, size_pt=10)
    r = p.add_run("e"); r.font.subscript = True; r.font.size = Pt(9)
    _run(p, "  +  w", color=BLACK, size_pt=10)
    r = p.add_run("u"); r.font.subscript = True; r.font.size = Pt(9)
    _run(p, " × SWT", color=BLACK, size_pt=10)
    r = p.add_run("u"); r.font.subscript = True; r.font.size = Pt(9)
    _run(p, "     (1)", color=BLACK, size_pt=10)

    body(doc, "where:")
    bullet(doc, "AWTₑ = average appointment waiting time for elective patients (hours), "
                "from call to scheduled appointment time; weight wₑ = 1/168 (max = 1 week = 168 h)")
    bullet(doc, "SWTᵤ = average scan waiting time for urgent patients (hours), "
                "from arrival to scan start; weight wᵤ = 1/9 (max = 1 working day = 9 h)")

    body(doc,
        "Four hypotheses guide the experiment. Main hypothesis: at least one tested combination "
        "of n_urgent, timing strategy, and appointment rule achieves a statistically significant "
        "reduction in the weighted objective W relative to the current baseline (n_urgent = 14, "
        "Strategy 1, Rule 1). Capacity hypothesis: a best intermediate value of n_urgent "
        "exists among the tested levels 10–20 that minimises W, because increasing n_urgent "
        "reduces SWTᵤ but raises AWTₑ. "
        "Timing hypothesis: Strategies 2 (uniform) and 3 (after-six) each achieve a lower W "
        "than Strategy 1 (end-of-block) through reduced SWTᵤ; the relative ordering of "
        "Strategies 2 and 3 is investigated empirically. "
        "Appointment-rule hypothesis: the appointment rule is expected to mainly affect "
        "elective scan waiting time and overtime, while having limited effect on the primary "
        "objective W. "
        "The experimental test design is described in Sections 3–7.")

    doc.add_paragraph()
    body(doc,
        "Three design decisions are examined: (1) the number of urgent slots per week, "
        "(2) the timing strategy for placing those slots within the weekly schedule, and "
        "(3) the appointment scheduling rule applied to elective patients. A full factorial "
        "discrete-event simulation experiment evaluates all 132 combinations of these factors "
        "to identify configurations that significantly outperform the current baseline. "
        "Discrete-event simulation (DES) is the standard methodology for evaluating "
        "such appointment scheduling systems, as it captures stochastic variability "
        "and time-dependent interactions that analytical models cannot (Law, 2015).")

    doc.add_page_break()

    # ══ 2. INPUT DATA ════════════════════════════════════════════════════════
    heading1(doc, "2.  Operational variability of the input data")

    body(doc,
        "The system exhibits significant operational variability from two sources: "
        "patient arrival processes and scan duration distributions.")

    heading2(doc, "Elective patient arrivals")
    body(doc,
        "Elective patients call the department on working days (Monday–Friday) between "
        "08:00 and 17:00. The daily number of calls follows a Poisson distribution with "
        "mean λₑ = 28.345 patients per day (negative-exponentially distributed "
        "inter-call times). On their appointment day, patients experience tardiness drawn "
        "independently from N(μ = 0, σ = 2.5) minutes. A no-show probability of 2% "
        "is applied independently per patient.")

    heading2(doc, "Urgent patient arrivals")
    body(doc,
        "Urgent patients arrive following a Poisson process with negative-exponentially "
        "distributed inter-arrival times. On full days (Monday, Tuesday, Wednesday, Friday) "
        "the rate is λᵤ = 2.5 patients/day; on half-days (Thursday, Saturday) "
        "λᵤ = 1.25 patients/day. No urgent arrivals occur outside opening hours, "
        "on Sundays, or during the lunch break.")

    heading2(doc, "Scan durations")
    body(doc,
        "Elective scan durations follow a truncated normal distribution "
        "N(μ = 15 min, σ = 3 min), truncated below at zero. "
        "Urgent scan durations depend on the scan type drawn from a discrete mixture "
        "distribution (Table 1). The truncated normal is used as a physical feasibility "
        "safeguard to exclude negative scan durations, which cannot occur in practice.")

    # Table 1
    t1 = doc.add_table(rows=6, cols=4)
    t1.style = 'Table Grid'
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(["Scan Type", "Frequency", "Mean (min)", "Std Dev (min)"]):
        tbl_header_cell(t1.rows[0].cells[c], h)
    for r, row in enumerate([
        ("Brain", "70 %", "15", "2.5"),
        ("Spine – Lumbar", "10 %", "17.5", "1.0"),
        ("Spine – Cervical", "10 %", "22.5", "2.5"),
        ("Abdomen MRCP", "5 %", "30", "1.0"),
        ("Others", "5 %", "30", "4.5"),
    ], 1):
        for c, v in enumerate(row):
            tbl_cell(t1.rows[r].cells[c], v)
    caption(doc, "Table 1. Urgent patient scan duration distributions (discrete mixture)")

    # ══ 3. RESPONSES ═════════════════════════════════════════════════════════
    heading1(doc, "3.  Evaluation output measure (responses)")

    body(doc, "Four key performance indicators (KPIs) are tracked per replication:")

    bullet(doc,
        "AWTₑ: average appointment waiting time for elective patients (hours). "
        "Reflects the time from call to scheduled appointment time; long waits indicate "
        "insufficient elective slot capacity relative to demand.",
        bold_prefix=None)
    bullet(doc,
        "SWTᵤ: average scan waiting time for urgent patients, from arrival to scan start, "
        "reported in minutes in all tables. For computation of W, SWTᵤ is converted to hours "
        "(÷60) and divided by the management threshold of 9 hours.",
        bold_prefix=None)
    bullet(doc,
        "SWTₑ: average scan waiting time for elective patients (minutes), from "
        "actual arrival (appointment time ± tardiness) to scan start. Mainly "
        "influenced by the appointment scheduling rule.",
        bold_prefix=None)
    bullet(doc,
        "OT: average daily overtime (minutes per open day), defined as the positive "
        "excess of the last scan end time over the scheduled closing time.",
        bold_prefix=None)

    doc.add_paragraph()
    body(doc,
        "The primary response is the weighted objective W from Equation (1). Both "
        "components are dimensionless: AWTₑ (hours) is divided by 168 and "
        "SWTᵤ (converted to hours) is divided by 9, normalising both relative to the "
        "management thresholds and ensuring balanced weighting. The components are not "
        "necessarily bounded by 1 because realised waiting times may exceed these thresholds.")

    # ══ 4. FACTORS ═══════════════════════════════════════════════════════════
    heading1(doc, "4.  Design parameters and values under consideration (factors)")

    body(doc,
        "Three design factors are investigated in a full factorial experiment of "
        "11 × 3 × 4 = 132 configurations.")

    heading2(doc, "Factor A – Number of urgent slots per week (n_urgent)")
    body(doc,
        "Total weekly slot capacity is fixed at 160. The number of urgent slots "
        "takes eleven values: all integers in the range 10–20 "
        "(10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20).")
    body(doc,
        "The remainder are elective slots. "
        "The current baseline uses 14 urgent slots (146 elective). "
        "Fewer urgent slots free more elective capacity (lower AWTₑ) but "
        "reduce the buffer for urgent demand (higher SWTᵤ), and vice versa.")

    heading2(doc, "Factor B – Timing strategy for urgent slots")
    body(doc, "Three strategies determine where urgent slots are placed in the weekly schedule:")
    bullet(doc, "Strategy 1 (End-of-block): urgent slots placed at the end of each "
                "morning/afternoon session block, distributed evenly across the 10 weekly blocks.")
    bullet(doc, "Strategy 2 (Uniform): urgent slots spread across each day using "
                "an even-spacing algorithm that minimises the maximum gap between "
                "consecutive urgent slots within each operating day.")
    bullet(doc, "Strategy 3 (After-six): within each operating day independently, one urgent "
                "slot is inserted after every 6 consecutive elective slots. The counter resets "
                "at each day boundary. Urgent slots per day are allocated proportionally to "
                "day length using Hamilton's largest-remainder method.")

    heading2(doc, "Factor C – Appointment scheduling rule for elective patients")
    body(doc,
        "Four rules relate the appointment time to the assigned slot start time. "
        "The Bailey-Welch rule and block-scheduling variants are well-established "
        "strategies for reducing patient waiting and scanner idle time in outpatient "
        "settings (Sickinger & Kolisch, 2009; Klassen & Rohleder, 1996):")
    bullet(doc, "Rule 1 (FCFS): appointment time = slot start time.")
    bullet(doc, "Rule 2 (Bailey-Welch, K=2): first two patients of each session "
                "scheduled at session start; subsequent patients at slot start − 15 min. "
                "Each morning block and each afternoon block constitutes a separate session, "
                "so the K=2 offset resets at 08:00 and again at 13:00 each full day.")
    bullet(doc, "Rule 3 (Blocking, B=2): patients sharing a pair of consecutive slots "
                "both scheduled at the first slot's start time.")
    bullet(doc, "Rule 4 (Benchmark, k = 0.5): appointment time = slot start − 1.5 min "
                "(= k × σₑ).")

    doc.add_paragraph()
    # Factor summary table
    t2 = doc.add_table(rows=4, cols=3)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(["Factor", "Name", "Levels"]):
        tbl_header_cell(t2.rows[0].cells[c], h)
    for r, row in enumerate([
        ("A", "Number of urgent slots (n_urgent)", "10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20"),
        ("B", "Timing strategy",
         "1 (end-of-block), 2 (uniform), 3 (after-six)"),
        ("C", "Appointment scheduling rule",
         "1 (FCFS), 2 (Bailey-Welch), 3 (Blocking), 4 (Benchmark)"),
    ], 1):
        for c, v in enumerate(row):
            tbl_cell(t2.rows[r].cells[c], v)
    caption(doc, "Table 2. Experimental factors and their levels (full factorial: 132 configurations)")

    doc.add_page_break()

    # ══ 5. SIMULATION DESIGN ═════════════════════════════════════════════════
    heading1(doc, "5.  The simulation design")

    heading2(doc, "5.1  State descriptor")
    body(doc,
        "The state descriptor contains all information needed to fully characterise "
        "the system at any simulation time (Law, 2015; Banks et al., 2010). "
        "It includes:")
    bullet(doc, "Current simulation time (minutes from Monday 08:00 of week 0)")
    bullet(doc, "Scanner status: idle or busy, and the scheduled end time of the current scan")
    bullet(doc, "Queue of elective patients waiting to access the scanner")
    bullet(doc, "Queue of urgent patients waiting to access the scanner")
    bullet(doc, "Remaining elective and urgent slots available on the current day")
    bullet(doc, "Current position in the cyclic schedule: week number, day (0–5), slot index")
    bullet(doc, "For each elective patient: call time, assigned slot, scheduled appointment time")
    bullet(doc, "For each urgent patient: arrival time, assigned urgent slot or overtime flag")
    doc.add_paragraph()
    body(doc,
        "Because appointment times are pre-assigned at call time, the state must maintain "
        "a forward-looking list of patients with pre-determined arrival targets. The simulation "
        "uses SimPy's process-based framework: each patient is modelled as an independent "
        "coroutine that sleeps until its appointment time, then requests the shared scanner.")

    heading2(doc, "5.2  Events")
    body(doc,
        "The simulation is implemented with SimPy (process-based DES library for Python). "
        "In a process-based DES framework, each entity is modelled as a concurrent process "
        "that advances through the simulation by yielding control at events such as "
        "arrivals, service requests, and completions (Law, 2015; Banks et al., 2010). "
        "The following events and processes drive the model:")
    bullet(doc, "ElectiveCallProcess (Mon–Fri, 08:00–17:00): assigns the next available "
                "elective slot and computes the appointment time using the active rule. "
                "Spawns an arrival process that fires at the appointment time.")
    bullet(doc, "ElectiveArrivalEvent: fires at appointment time with tardiness from "
                "N(0, 2.5²) min; 2% no-show check applied. Showing patients join the queue.")
    bullet(doc, "UrgentArrivalProcess: active during opening hours only (morning and afternoon "
                "blocks; no arrivals during lunch, outside opening hours, Sundays, or half-day "
                "afternoons). Inter-arrivals exponentially distributed. Each urgent patient is "
                "assigned to the earliest remaining urgent slot on the same day with a start time "
                "at or after the patient's arrival time. If no same-day urgent slot remains, the "
                "patient is assigned to overtime after the regular closing time of that day, after "
                "all scheduled patients have been served.")
    body(doc,
        "Queue discipline: Both urgent and elective patients request the scanner via a shared "
        "SimPy Resource queue (FIFO discipline within each time slot). Priority is established "
        "structurally through the slot assignment mechanism: urgent patients wait until their "
        "assigned urgent slot time and then compete for the scanner alongside any elective "
        "patients whose appointment time has also arrived. No preemptive priority for urgent "
        "patients over elective patients is implemented in the model.")
    bullet(doc, "ServiceProcess: scanner modelled as a SimPy Resource (capacity 1). "
                "Scan duration drawn from the appropriate distribution; next queued patient "
                "begins service on completion.")
    bullet(doc, "DayStartEvent / DayEndEvent: reset slot counters; record daily overtime.")
    bullet(doc, "WeekStartEvent: advances cyclic schedule position (pattern repeats identically).")

    heading2(doc, "5.3  Agenda")
    body(doc,
        "Table 3 illustrates a representative event sequence for an urgent Brain MRI patient "
        "arriving on a Monday morning.")

    t3 = doc.add_table(rows=5, cols=2)
    t3.style = 'Table Grid'
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(["Time", "Event"]):
        tbl_header_cell(t3.rows[0].cells[c], h)
    for r, (tim, evt) in enumerate([
        ("08:47", "Urgent arrival – assigned to urgent slot 3 (start: 10:15)"),
        ("10:15", "Slot time reached – patient enters scanner queue"),
        ("10:23", "Scanner idle – patient begins service (Brain scan, 15 min)"),
        ("10:38", "Service completion – scanner becomes idle"),
    ], 1):
        tbl_cell(t3.rows[r].cells[0], tim)
        tbl_cell(t3.rows[r].cells[1], evt)
    caption(doc, "Table 3. Example event agenda for an urgent Brain MRI patient")

    heading2(doc, "5.4  Main routine")
    body(doc,
        "The main routine initialises the SimPy environment, random streams, the scanner "
        "resource, and the weekly slot schedule. It spawns all recurring processes and runs "
        "the environment for (warm-up weeks + 52 weeks) × 10,080 minutes. Data collection "
        "begins only after the warm-up cutoff. The stopping criterion is time-based, "
        "consistent with the cyclic weekly structure.")
    body(doc,
        "Key implementation assumptions: (1) Patient tardiness is drawn from N(0, 2.5²) min; "
        "negative values indicate early arrival and are retained as-is. "
        "(2) Overtime is measured as the positive excess of the last scan completion time over "
        "the scheduled day-end time (17:00 on full days, 12:00 on half-days); it is zero if "
        "all scans finish on time. "
        "(3) The scanner queue within each slot uses SimPy's FIFO Resource discipline; "
        "no preemptive priority is implemented. "
        "(4) Urgent patients who arrive when all same-day urgent slots are taken are assigned "
        "to an overtime slot that begins after all regularly scheduled patients have been "
        "served on that day. "
        "(5) Scans that begin before 12:00 may extend into the lunch window without "
        "interruption; only the scan start is blocked. "
        "(6) No scan may start during the lunch break (12:00–13:00 on full days). If a "
        "patient acquires the scanner during this interval, the scan start is delayed to "
        "13:00; the scanner resource is held throughout the wait, preserving FIFO order.")

    heading2(doc, "5.5  Simulation time")
    body(doc,
        "Each replication simulates 72 virtual weeks: 20 weeks warm-up plus 52 weeks "
        "(one full year) of data collection. One week = 7 × 1,440 = 10,080 min. "
        "Data before the warm-up cutoff (minute 201,600) is discarded. "
        "The full factorial experiment (132 configurations × 57 replications = 7,524 runs) "
        "ran for approximately 273 wall-clock minutes using 8 CPU cores "
        "(Python ProcessPoolExecutor).")

    doc.add_page_break()

    # ══ 6. HYPOTHESES ════════════════════════════════════════════════════════
    heading1(doc, "6.  Hypotheses")

    body(doc,
        "The experiment tests the following main hypothesis and three subhypotheses. "
        "Hypotheses are stated in null/alternative form to support formal statistical testing.")

    p = body_para(doc)
    _run(p, "Main hypothesis: ", bold=True, color=BLACK, size_pt=10)
    _run(p, "H₀: no tested configuration has a lower mean W than the baseline "
            "(n_urgent = 14, Strategy 1, Rule 1). "
            "H₁: at least one tested non-baseline configuration has a lower mean W than the baseline. "
            "The experiment uses paired CRN t-tests with Bonferroni correction to test H₁.",
        color=BLACK, size_pt=10)

    doc.add_paragraph()
    body(doc, "Three subhypotheses address the mechanisms driving any improvement:")

    bullet(doc,
        "Capacity hypothesis: Increasing n_urgent is expected to reduce urgent scan waiting "
        "time (SWTᵤ) because more urgent slots are available per day, but to increase "
        "appointment waiting time for elective patients (AWTₑ) because fewer elective slots "
        "are available. Consequently, W may show a trade-off or U-shaped pattern across the "
        "tested urgent-slot levels 10–20. No formal claim of convexity is "
        "made without empirical verification.",
        bold_prefix=None)
    bullet(doc,
        "Timing strategy hypothesis: Strategies 2 (uniform spacing) and 3 (after-six) are "
        "each expected to reduce SWTᵤ and W relative to Strategy 1 (end-of-block) because "
        "both distribute urgent slots more evenly across operating hours, reducing the maximum "
        "time an urgent patient must wait for the next available slot. The relative ordering "
        "of Strategies 2 and 3 is an empirical question investigated in this experiment.",
        bold_prefix=None)
    bullet(doc,
        "Appointment-rule hypothesis: The appointment scheduling rule (Factor C) is expected "
        "to have limited impact on the primary objective W, because W is driven by AWTₑ and "
        "SWTᵤ — neither of which is directly controlled by the appointment-time offset. "
        "However, the rule may affect elective scan waiting time (SWTₑ) and overtime (OT). "
        "Therefore, appointment-rule choice is treated as a secondary decision, serving as "
        "a tie-breaker when W differences across rules are statistically or practically small.",
        bold_prefix=None)

    doc.add_page_break()

    # ══ 7. EXPERIMENTAL DESIGN ═══════════════════════════════════════════════
    heading1(doc, "7.  Choice of experimental design")

    heading2(doc, "7.1  Full factorial design")
    body(doc,
        "A complete 11 × 3 × 4 full factorial design was chosen because 132 "
        "configurations are computationally manageable and the full factorial allows "
        "unconfounded estimation of all main effects and two-way interactions "
        "(Montgomery, 2017). "
        "A fractional design would reduce computation but alias interaction effects "
        "that may be operationally relevant; the full factorial avoids this trade-off "
        "(Law, 2015; Montgomery, 2017).")

    heading2(doc, "7.2  Variance reduction techniques – Common Random Numbers (CRN)")
    body(doc,
        "Common Random Numbers (CRN) is a variance reduction technique that improves the "
        "precision of pairwise comparisons by inducing positive correlation between "
        "simulation outputs across configurations driven by the same random variates "
        "(Law, 2015; Montgomery, 2017). "
        "Each replication index is assigned six independent random streams: "
        "elective call inter-arrivals, urgent inter-arrivals, elective scan durations, "
        "urgent scan durations, patient tardiness, and no-show decisions. Seeds are "
        "derived deterministically via SHA-256 hashing of the stream name and replication "
        "index, ensuring full reproducibility across all 132 configurations.")
    body(doc,
        "Because all configurations share the same sequence of random variates for each "
        "replication, differences in outcomes reflect structural configuration differences "
        "rather than sampling noise. This substantially reduces the variance of paired "
        "differences, which is central to the paired t-test comparisons (Law, 2015).")

    heading2(doc, "7.3  Warm-up period")
    body(doc,
        "The warm-up period was determined using Welch's method (Welch, 1983; Law, 2015), "
        "which identifies the end of the initial transient "
        "by inspecting a smoothed moving average of the output process across replications. "
        "Ten pilot replications of the baseline configuration were simulated for 72 weeks. "
        "The weekly objective W was averaged across replications and smoothed with a "
        "moving average of window w = 4 weeks:")

    p = body_para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "ȳᵢ(w) = (2w+1)⁻¹ × Σ_{j=i−w}^{i+w} Ȳⱼ,   for i = w+1, ..., T−w    (2)",
         color=BLACK, size_pt=10)

    body(doc,
        "where Ȳⱼ denotes the average of the weekly objective W across all pilot "
        "replications at simulation week j.")

    doc.add_paragraph()
    body(doc,
        "Visual inspection of the moving-average plot (Figure 1) identified stabilisation "
        "after approximately 16–20 weeks. A conservative warm-up of 20 weeks was adopted. "
        "The high server utilisation at baseline (estimated ρ ≈ 0.97 based on scan demand "
        "relative to available capacity) causes slow convergence; "
        "the 20-week threshold removes the initial transient while acknowledging the "
        "system's inherent week-to-week variability. "
        "Configurations with n_urgent ≥ 18 are expected to result in server utilisation "
        "ρ ≥ 1.0 based on analytical capacity estimates (n=19 and n=20 leave fewer elective "
        "slots than expected weekly elective demand; n=18 is borderline). Under such conditions "
        "no steady state exists and the elective appointment queue grows without bound; "
        "the simulation estimates do not represent stable long-run performance. "
        "These configurations are excluded from all recommendations. "
        "Confidence intervals reported for these configurations are not meaningful.")

    fig(doc, "warmup_plot.png", w=5.5,
        cap="Figure 1. Welch warm-up plot for the baseline configuration based on 10 pilot "
            "replications. The curve shows the moving average of weekly W with smoothing "
            "half-window w = 4; data collection starts after a conservative 20-week warm-up.")

    heading2(doc, "7.4  Number of simulation runs")
    body(doc,
        "The required number of replications was determined using the pilot run method "
        "(Law, 2015; Montgomery, 2017). "
        "The sample standard deviation of the per-replication objective W, computed from "
        "10 pilot replications, was used to derive the minimum n for relative precision "
        "ε = 5% at α = 0.05:")

    p = body_para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "n  ≥  ceil( ( t₀.₉₇₅, n₀₋₁  ×  s₀  /  (ε × W̄₀) )² )    (3)",
         color=BLACK, size_pt=10)

    doc.add_paragraph()
    body(doc,
        "Based on n₀ = 10 pilot replications of the baseline configuration, the sample "
        "mean W̄₀ = 0.442 and sample standard deviation s₀ = 0.074 were obtained. "
        "With t₀.₉₇₅,₉ = 2.262 (t-distribution, df = n₀ − 1 = 9) and relative precision "
        "target ε = 0.05, the required number of replications is:")

    p = body_para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "n = ceil( (2.262 × 0.074 / (0.05 × 0.442))² ) = ceil(56.8) = 57",
         color=BLACK, size_pt=10)

    doc.add_paragraph()
    body(doc,
        "The full factorial experiment used 57 replications per configuration, satisfying "
        "the pilot-based criterion. The achieved relative precision with 57 replications "
        "is 3.23% for the baseline W (95% CI half-width = 0.014; W̄ = 0.438), well "
        "within the 5% target. The pilot standard deviation (s₀ = 0.074) slightly "
        "overestimated the true variability; the actual standard deviation is 0.053. "
        "Results are reported with achieved confidence intervals as the primary precision "
        "indicator.")

    heading2(doc, "7.5  Statistical analysis")
    body(doc, "The following procedures were applied to the simulation output:")
    bullet(doc,
        "95% confidence intervals for the mean of W and all KPIs per configuration, "
        "using the t-distribution with n−1 = 56 degrees of freedom (Law, 2015).")
    bullet(doc,
        "Blocked factorial ANOVA on replication-level output, with the replication index "
        "as a blocking factor (model: W ~ C(n_urgent) × C(strategy) × C(rule) + C(rep_id)). "
        "Including the replication block accounts for the CRN structure and gives unbiased "
        "estimates of factor effects (Montgomery, 2017). Main effects and all two-way "
        "interactions were tested.")
    bullet(doc,
        "Paired CRN t-tests: each non-baseline configuration was compared to the baseline "
        "(n_urgent = 14, Strategy 1, Rule 1) using per-replication differences in W, "
        "exploiting the positive correlation induced by CRN (Law, 2015).")
    bullet(doc,
        "Planned CRN paired comparisons between configurations of primary interest "
        "(n=12 vs n=11 and n=13 at S3, R2; strategy comparison S3 vs S2 at n=12, R2), "
        "with Bonferroni correction for the number of planned comparisons within each "
        "group (Montgomery, 2017).")
    bullet(doc,
        "Familywise Bonferroni correction (×131) for the 131 vs-baseline comparisons, "
        "significance threshold αᴮ = 0.05/131 ≈ 0.000382 (Montgomery, 2017).")

    heading2(doc, "7.6  Results")
    body(doc,
        "The baseline configuration (n_urgent = 14, Strategy 1, Rule 1) achieves "
        "W = 0.438 (95% CI: [0.424, 0.452]), with AWTₑ = 24.3 hours, SWTᵤ = 158.1 minutes, "
        "and OT = 18.4 minutes per open day. "
        "Table 4 shows the top-10 configurations ranked by W.")

    # Top-10 table (updated with new results)
    top10_hdr = ["Rank", "n", "S", "R", "W", "95 % CI", "AWTₑ (h)",
                 "SWTᵤ (min)", "SWTₑ (min)", "OT (min/d)"]
    top10_rows = [
        ("1",  "12","3","2","0.333","[0.324, 0.341]","17.4","123.6","11.1","24.5"),
        ("2",  "13","3","2","0.334","[0.323, 0.344]","20.0","115.8","10.9","23.0"),
        ("3",  "12","3","4","0.334","[0.326, 0.343]","17.7","123.8","5.3","27.9"),
        ("4",  "13","3","4","0.334","[0.324, 0.345]","20.2","115.6","5.2","26.5"),
        ("5",  "12","3","3","0.335","[0.327, 0.344]","17.6","124.7","9.6","26.9"),
        ("6",  "13","3","3","0.336","[0.326, 0.346]","20.1","116.7","9.5","25.5"),
        ("7",  "12","3","1","0.337","[0.328, 0.345]","17.7","124.9","5.2","29.0"),
        ("8",  "11","3","2","0.337","[0.330, 0.343]","15.6","131.7","11.3","25.9"),
        ("9",  "13","3","1","0.337","[0.326, 0.347]","20.2","116.7","5.2","27.6"),
        ("10", "11","3","4","0.339","[0.332, 0.345]","15.8","132.2","5.3","29.2"),
    ]
    t4 = doc.add_table(rows=11, cols=10)
    t4.style = 'Table Grid'
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(top10_hdr):
        tbl_header_cell(t4.rows[0].cells[c], h, size=8)
    for r, row in enumerate(top10_rows, 1):
        for c, v in enumerate(row):
            tbl_cell(t4.rows[r].cells[c], v, size=8)
    caption(doc, "Table 4. Top-10 configurations by W (57 reps each). "
            "n = n_urgent, S = strategy, R = rule. Baseline (n=14, S1, R1): W = 0.438. "
            "SWTᵤ reported in minutes; converted to hours (÷60) for computation of W. "
            "SWTₑ column required for secondary-objective tie-breaking per assignment specification.")

    body(doc,
        "All ten top configurations use Strategy 3 (after-six placement), confirming the "
        "strong effect of timing strategy. Strategy 3 achieves the lowest mean SWTᵤ "
        "(110.3 min averaged across n and rule levels) compared to Strategy 2 (126.7 min) "
        "and Strategy 1 (151.0 min). The after-six pattern places urgent slots at regular "
        "intervals within each day (every 7th slot), creating shorter maximum waiting gaps "
        "for urgent patients than either uniform or end-of-block placement. "
        "The best tested configuration (n_urgent = 12, S3, R2) achieves W = 0.333 "
        "(95% CI: [0.324, 0.341]), a 24% reduction versus baseline "
        "(paired CRN t-test: Δ = −0.105, 95% CI [−0.111, −0.098], t(56) = −31.7, "
        "pᴮ < 0.001), statistically significant after Bonferroni correction. "
        "The improvement is driven by reducing SWTᵤ from 158.1 to 123.6 minutes (−22%) "
        "and AWTₑ from 24.3 to 17.4 hours (−28%). "
        "Note that CRN induces very high statistical power: even operationally negligible "
        "W differences (Δ < 0.005) yield highly significant p-values; practical significance "
        "should be judged from the magnitude of differences, not p-values alone.")

    body(doc,
        "Of the 131 non-baseline configurations, 66 significantly outperform the baseline "
        "(Bonferroni-corrected pᴮ < 0.000382), 50 are significantly worse, and 15 show no "
        "statistically significant difference. All significantly better configurations use "
        "either Strategy 2 or Strategy 3 with n_urgent ≤ 15. All significantly worse "
        "configurations have n_urgent ≥ 16, driven by elective slot shortage and the "
        "resulting growth in AWTₑ.")

    body(doc,
        "Blocked factorial ANOVA results (model: W ~ C(n_urgent)×C(strategy)×C(rule) + "
        "C(rep_id); R² = 0.786): n_urgent is the dominant factor (F(10, 7336) = 1985.4, "
        "p < 0.001, η² = 0.578). Strategy is significant (F(2, 7336) = 239.7, p < 0.001, "
        "η² = 0.014). The appointment scheduling rule has no statistically significant "
        "effect on W (F(3, 7336) = 0.400, p = 0.753). No two-way interaction is significant "
        "(n_urgent × strategy: F(20, 7336) = 0.371, p = 0.995). The CRN replication block "
        "accounts for 19.4% of total variance (F(56, 7336) = 119.1, p < 0.001), confirming "
        "that CRN induced substantial positive correlation across configurations and justified "
        "its use. Residual stochastic noise supports the use of confidence intervals and "
        "paired comparisons as primary inference tools. "
        "All configurations with n_urgent ≥ 16 perform significantly worse than the "
        "baseline due to elective backlog caused by reduced elective slot capacity.")

    heading2(doc, "7.6a  Optimal n_urgent: planned comparison")
    body(doc,
        "The top-10 table shows that n=12 and n=13 both appear frequently, while n=11 "
        "appears less. A planned CRN paired comparison at Strategy 3, Rule 2 was performed "
        "to determine the optimal n_urgent (Bonferroni adjustment ×3 for three comparisons):")

    # n comparison table
    cmp_hdr = ["Comparison", "ΔW (mean)", "95% CI", "t(56)", "p-value", "pᴮ (×3)"]
    cmp_rows = [
        ("n=12,S3,R2 vs n=11,S3,R2", "−0.0038", "[−0.0056, −0.0019]", "−4.13", "0.00012", "0.00037"),
        ("n=12,S3,R2 vs n=13,S3,R2", "−0.0007", "[−0.0029, +0.0016]", "−0.62",  "0.541",   "1.000"),
        ("n=12,S3,R2 vs n=10,S3,R2", "−0.0096", "[−0.0128, −0.0063]", "−5.92", "< 0.001",  "< 0.001"),
    ]
    t_cmp = doc.add_table(rows=4, cols=6)
    t_cmp.style = 'Table Grid'
    t_cmp.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(cmp_hdr):
        tbl_header_cell(t_cmp.rows[0].cells[c], h, size=8)
    for r, row in enumerate(cmp_rows, 1):
        for c, v in enumerate(row):
            tbl_cell(t_cmp.rows[r].cells[c], v, size=8)
    caption(doc, "Table 5. Planned CRN paired comparisons for n_urgent at S3, R2. "
            "ΔW = W(n=12) − W(comparison). Bonferroni adjustment ×3. "
            "n=12 and n=13 are statistically indistinguishable (pᴮ = 1.000).")

    body(doc,
        "n=12 is significantly better than n=11 (p = 0.00012, pᴮ = 0.00037) and n=10 "
        "(pᴮ < 0.001), but not significantly different from n=13 (p = 0.54). "
        "The W difference between n=12 and n=13 is negligible (Δ = 0.0007). "
        "However, there is an important trade-off: n=12 achieves AWTₑ = 17.4 h and "
        "SWTᵤ = 123.6 min, while n=13 achieves AWTₑ = 20.0 h and SWTᵤ = 115.8 min. "
        "Since the primary objective W is identical, n=12 is preferred because it delivers "
        "substantially shorter appointment waiting times for elective patients (AWTₑ 2.6 hours "
        "lower), at only a modest increase in SWTᵤ (7.8 minutes). "
        "This supports selecting n=12 as the recommended capacity level.")

    heading2(doc, "7.6b  Appointment-rule selection: secondary objectives")
    body(doc,
        "The blocked ANOVA confirms that appointment rule has no significant effect on W "
        "(p = 0.753). CRN-paired comparisons within n=12, S3 reveal statistically significant "
        "W differences between rules (Bonferroni ×3), but the magnitudes are small "
        "(max Δ(W) = 0.004). The secondary objectives SWTₑ and OT therefore support the "
        "appointment-rule recommendation:")

    # Rule comparison table at n=12, S3
    rule_hdr = ["Rule", "W (mean)", "SWTₑ (min)", "SWTᵤ (min)", "OT (min/d)", "ΔW vs R2*", "pᴮ (×3)†"]
    rule_rows = [
        ("R1 FCFS",          "0.337", "5.2",  "124.9", "29.0", "+0.0037", "< 0.001"),
        ("R2 Bailey-Welch",  "0.333", "11.1", "123.6", "24.5", "—",       "—"),
        ("R3 Blocking",      "0.335", "9.6",  "124.7", "26.9", "+0.0025", "< 0.001"),
        ("R4 Benchmark",     "0.334", "5.3",  "123.8", "27.9", "+0.0016", "< 0.001"),
    ]
    t_rule = doc.add_table(rows=5, cols=7)
    t_rule.style = 'Table Grid'
    t_rule.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(rule_hdr):
        tbl_header_cell(t_rule.rows[0].cells[c], h, size=8)
    for r, row in enumerate(rule_rows, 1):
        for c, v in enumerate(row):
            tbl_cell(t_rule.rows[r].cells[c], v, size=8)
    caption(doc, "Table 6. Appointment rules at n=12, Strategy 3 (57 reps, CRN). "
            "* ΔW = W(rule) − W(R2); † Bonferroni-adjusted p-value (×3). "
            "R2 achieves the lowest W, lowest SWTᵤ, and lowest OT. "
            "R1 and R4 achieve lower SWTₑ at the cost of higher W and OT.")

    body(doc,
        "Rule 2 (Bailey-Welch) achieves the lowest W (0.333) and also the lowest overtime "
        "(24.5 min/day) and lowest SWTᵤ (123.6 min). The W differences across rules are "
        "statistically significant after Bonferroni correction but operationally small "
        "(max Δ = 0.004; ANOVA: F(3, 7336) = 0.400, p = 0.753). The drawback of Rule 2 "
        "is higher SWTₑ (11.1 min) versus Rules 1 and 4 (~5.2 min), because Bailey-Welch "
        "deliberately schedules early patients at session start, creating intentional "
        "overlap. If minimising elective scan queue time is the department's secondary "
        "priority, Rule 1 (FCFS) or Rule 4 (Benchmark) are preferable alternatives. "
        "The recommended configuration is therefore n_urgent = 12, Strategy 3, Rule 2. "
        "Rule 1 or Rule 4 are acceptable alternatives if minimising SWTₑ is the "
        "department's secondary priority.")

    fig(doc, "main_effects_plot.png", w=5.5,
        cap="Figure 2. Main effects of n_urgent, strategy, and appointment rule on W")
    fig(doc, "lines_per_strategy.png", w=5.5,
        cap="Figure 3. Mean W by n_urgent for each timing strategy (with 95 % CI bands)")
    fig(doc, "top10_bar.png", w=5.0,
        cap="Figure 4. Top-10 configurations: mean W with 95 % CI and baseline reference")

    doc.add_page_break()

    # ══ 8. CONCLUSION ════════════════════════════════════════════════════════
    heading1(doc, "8.  Conclusion")

    body(doc,
        "This study investigated the best appointment scheduling configuration for a "
        "single-server outpatient radiology department using a discrete-event simulation "
        "model. A full factorial experiment over 132 configurations with 57 replications "
        "per configuration provided performance estimates meeting the pilot-formula precision "
        "criterion (achieved relative precision 3.23%, within the 5% target). "
        "Common Random Numbers ensured low-variance pairwise comparisons, and a "
        "20-week Welch warm-up removed initial transient bias.")

    body(doc,
        "Among the 132 tested configurations, the lowest estimated W is achieved by "
        "n_urgent = 12 with after-six spacing (Strategy 3) and Bailey-Welch scheduling "
        "(Rule 2), with W = 0.333 (95% CI: [0.324, 0.341]). This represents a 24% reduction "
        "versus the baseline (W = 0.438), statistically significant after Bonferroni correction "
        "(Δ = −0.105, pᴮ < 0.001). The improvement is driven by reducing both AWTₑ (24.3 → "
        "17.4 hours, −28%) and SWTᵤ (158.1 → 123.6 minutes, −22%). "
        "The appointment scheduling rule does not significantly affect W (ANOVA: "
        "F(3, 7336) = 0.400, p = 0.753). Rule 2 (Bailey-Welch) is recommended because it "
        "also achieves the lowest overtime (24.5 min/day). Rule 1 (FCFS) or Rule 4 "
        "(Benchmark) are acceptable alternatives if minimising SWTₑ (~5.2 min vs 11.1 min "
        "for Rule 2) is the department's secondary priority. "
        "The recommended configuration is n_urgent = 12, Strategy 3, Rule 2.")

    body(doc,
        "Three conclusions follow from the experiment. "
        "First, n_urgent is the dominant factor (η² = 0.578): the current baseline allocates "
        "14 urgent slots, which is two more than the best tested level. Reducing to 12 urgent "
        "slots improves AWTₑ from 24.3 to 17.4 hours while keeping SWTᵤ at 123.6 minutes — "
        "a favourable trade-off. n=12 and n=13 are statistically indistinguishable on W "
        "(Δ = 0.0007, p = 0.54), but n=12 is preferred because it delivers 2.6 hours lower "
        "AWTₑ with only 7.8 minutes higher SWTᵤ. "
        "Second, after-six spacing (Strategy 3) outperforms both uniform (Strategy 2) and "
        "end-of-block (Strategy 1) placement (η² = 0.014; Strategy 3 mean W = 0.489 vs "
        "0.520 for S2 and 0.568 for S1). The after-six pattern creates shorter maximum "
        "intra-day gaps between urgent slots, reducing SWTᵤ by 16.4 minutes on average "
        "compared to uniform spacing and 40.7 minutes compared to end-of-block. "
        "Third, the appointment scheduling rule does not affect the primary objective W "
        "(η² ≈ 0), confirming the appointment-rule hypothesis.")

    body(doc,
        "Regarding hypothesis outcomes: the main hypothesis is confirmed (66 of 131 "
        "non-baseline configurations are significantly better after Bonferroni correction). "
        "The capacity hypothesis is confirmed: W decreases as n_urgent decreases from 20 to "
        "12, but increases again for n_urgent < 12 (n=11 is significantly worse than n=12), "
        "identifying n=12 as the optimal level. The timing strategy hypothesis is confirmed "
        "in direction (Strategies 2 and 3 both outperform Strategy 1), with the additional "
        "finding that Strategy 3 outperforms Strategy 2 — an outcome not anticipated in "
        "the original hypothesis. The appointment-rule hypothesis is confirmed.")

    body(doc,
        "The main limitation is the high estimated server utilisation (ρ ≈ 0.97 at baseline), "
        "which causes persistent oscillations and makes the warm-up cutoff approximate. "
        "Configurations with n_urgent ≥ 18 appear to approach or exceed ρ = 1.0 "
        "and are excluded from recommendations. Future work could explore demand smoothing, "
        "extended operating hours, or priority queue disciplines for urgent patients.")

    doc.add_page_break()

    # ══ REFERENCES ═══════════════════════════════════════════════════════════
    p = _para(doc, before=0, after=8)
    _run(p, "References", bold=True, color=BLUE, size_pt=13)

    refs = [
        ("Banks, J., Carson, J. S., Nelson, B. L., & Nicol, D. M. (2010). "
         "Discrete-Event System Simulation (5th ed.). Pearson Prentice Hall."),
        ("Klassen, K. J., & Rohleder, T. R. (1996). Scheduling outpatient appointments in a "
         "dynamic environment. Journal of Operations Management, 14(2), 83–101."),
        ("Law, A. M. (2015). Simulation Modeling and Analysis (5th ed.). McGraw-Hill Education."),
        ("Montgomery, D. C. (2017). Design and Analysis of Experiments (9th ed.). Wiley."),
        ("Sickinger, S., & Kolisch, R. (2009). The performance of a generalized Bailey-Welch "
         "rule for outpatient appointment scheduling under inpatient and emergency demand. "
         "Health Care Management Science, 12(4), 408–419."),
        ("Team SimPy. (2023). SimPy: Discrete event simulation for Python (Version 4). "
         "https://simpy.readthedocs.io/"),
        ("Virtanen, P., et al. (2020). SciPy 1.0: Fundamental algorithms for scientific "
         "computing in Python. Nature Methods, 17, 261–272. "
         "https://doi.org/10.1038/s41592-020-0772-5"),
        ("Welch, P. D. (1983). The statistical analysis of simulation results. "
         "In S. S. Lavenberg (Ed.), Computer Performance Modeling Handbook "
         "(pp. 268–328). Academic Press."),
    ]
    for ref in refs:
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=2, after=4)
        p.paragraph_format.left_indent     = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        _run(p, ref, color=BLACK, size_pt=10)

    doc.add_page_break()

    # ══ APPENDIX ═════════════════════════════════════════════════════════════
    heading_appendix(doc, "Appendix A – Additional plots")

    fig(doc, "pareto_tradeoff.png", w=5.5,
        cap="Figure A1. Pareto trade-off between AWTₑ and SWTᵤ across all 132 configurations")
    fig(doc, "interaction_plot.png", w=5.5,
        cap="Figure A2. Interaction effects between n_urgent and strategy on W")
    fig(doc, "heatmap_objective.png", w=5.5,
        cap="Figure A3. Heatmap of mean W for n_urgent × strategy combinations")
    fig(doc, "rule_boxplot.png", w=5.5,
        cap="Figure A4. Distribution of W, AWTₑ and SWTᵤ by appointment rule")
    fig(doc, "kpi_breakdown.png", w=5.5,
        cap="Figure A5. KPI contribution: AWTₑ vs SWTᵤ components of W per n_urgent level")

    doc.save(OUT)
    print(f"Report saved: {OUT}")

if __name__ == "__main__":
    build()
